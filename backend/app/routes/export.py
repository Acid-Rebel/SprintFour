"""
Export route — produces a safe document with PII replaced by tokens,
plus a safety report proving irreversibility.
"""

from fastapi import APIRouter, UploadFile, File, Form
from fastapi.responses import Response
import json
import fitz  # PyMuPDF

from ..models.schemas import ExportRequest, ExportResponse

router = APIRouter(prefix="/api", tags=["export"])


@router.post("/export", response_model=ExportResponse)
async def export_safe_document(req: ExportRequest):
    """Replace all redacted spans with type-labelled tokens and return a safety report."""

    text = req.document_text
    # Sort redactions by start position in reverse so replacements don't shift indices
    active = sorted(
        [r for r in req.redactions if r.is_redacted],
        key=lambda r: r.start,
        reverse=True,
    )

    redacted_items = []
    for r in active:
        token = f"[{r.pii_type}]"
        text = text[:r.start] + token + text[r.end:]
        redacted_items.append({
            "type": r.pii_type,
            "confidence": r.confidence,
            "was_user_override": r.user_override,
        })

    overrides = [r for r in req.redactions if r.user_override]
    kept_by_user = [r for r in req.redactions if not r.is_redacted]

    safety_report = {
        "total_pii_found": len(req.redactions),
        "total_redacted": len(active),
        "total_kept_by_user": len(kept_by_user),
        "user_overrides": len(overrides),
        "redacted_items": redacted_items,
        "irreversibility_proof": (
            "All redacted text has been replaced with type-labelled placeholder "
            "tokens (e.g. [NAME], [EMAIL]). The original text is NOT embedded, "
            "encoded, or hidden anywhere in the exported output. The replacement "
            "is a destructive string substitution — the original characters no "
            "longer exist in this document."
        ),
        "remaining_risks": (
            "Even with PII removed, surrounding context may allow "
            "re-identification in some cases. Review context-risk warnings."
        ),
    }

    return ExportResponse(safe_text=text, safety_report=safety_report)

@router.post("/export-pdf")
async def export_pdf_document(
    file: UploadFile = File(...),
    redaction_boxes: str = Form(...) # JSON string of boxes
):
    """Truly redacts a PDF while maintaining vector quality by removing text from the content stream."""
    pdf_bytes = await file.read()
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    
    boxes = json.loads(redaction_boxes)
    
    for box in boxes:
        page_index = box.get("pageNum") - 1
        if page_index < 0 or page_index >= len(doc):
            continue
            
        page = doc[page_index]
        page_height = page.rect.height
        
        # PDF.js provides unscaled origin at bottom-left: tx, ty
        # PyMuPDF uses top-left origin.
        tx = box.get("rawX")
        ty = box.get("rawY")
        w = box.get("rawWidth")
        h = box.get("rawHeight")
        
        x0 = tx
        y1 = page_height - ty + (h * 0.2) # Bottom of text + 20% for descenders
        y0 = y1 - (h * 1.2) # Top of text + 20% for ascenders
        x1 = x0 + w
        
        rect = fitz.Rect(x0, y0, x1, y1)
        
        # Add redaction annotation (removes text underneath and draws a black box)
        page.add_redact_annot(rect, fill=(0, 0, 0))
        
    # Apply all redaction annotations on all pages
    for page in doc:
        page.apply_redactions()
        
    out_bytes = doc.write()
    doc.close()
    
    return Response(
        content=out_bytes, 
        media_type="application/pdf", 
        headers={"Content-Disposition": "attachment; filename=redacted.pdf"}
    )

@router.post("/export-docx")
async def export_docx_document(
    file: UploadFile = File(...),
    redactions: str = Form(...) # JSON string of redactions
):
    """Edits a DOCX file in-place by replacing text with tokens."""
    import io
    from docx import Document
    
    doc_bytes = await file.read()
    doc = Document(io.BytesIO(doc_bytes))
    
    redaction_data = json.loads(redactions)
    visible_redactions = [r for r in redaction_data if r.get('is_redacted')]
    
    # Simple search and replace across paragraphs and table cells
    # This replaces the entire text of the run to preserve as much formatting as possible,
    # but in python-docx, modifying paragraph.text strips inline styles.
    # To do it properly while keeping styles, we replace text inside the individual runs.
    
    def replace_in_paragraph(p):
        for r in visible_redactions:
            orig = r.get('original_text')
            pii_type = f"[{r.get('pii_type')}]"
            
            while orig in p.text:
                full_text = p.text
                start_idx = full_text.find(orig)
                end_idx = start_idx + len(orig)
                
                curr_idx = 0
                first_run_found = False
                
                for run in p.runs:
                    run_len = len(run.text)
                    run_start = curr_idx
                    run_end = curr_idx + run_len
                    
                    if run_end > start_idx and run_start < end_idx:
                        # This run intersects with the original text
                        if run_start <= start_idx and run_end >= end_idx:
                            # Completely contains the text
                            run.text = run.text[:start_idx - run_start] + pii_type + run.text[end_idx - run_start:]
                            break
                            
                        if not first_run_found:
                            # First run that intersects
                            run.text = run.text[:start_idx - run_start] + pii_type
                            first_run_found = True
                        else:
                            # Subsequent intersecting runs
                            if run_end >= end_idx:
                                # Last run that intersects
                                run.text = run.text[end_idx - run_start:]
                            else:
                                # Completely inside the text
                                run.text = ""
                                
                    curr_idx += run_len
                    
    for p in doc.paragraphs:
        replace_in_paragraph(p)
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)
                    
    out_io = io.BytesIO()
    doc.save(out_io)
    out_bytes = out_io.getvalue()
    
    return Response(
        content=out_bytes, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
        headers={"Content-Disposition": "attachment; filename=redacted.docx"}
    )

